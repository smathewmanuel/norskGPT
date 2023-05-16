import os
import openai
import yaml
from docx import Document


def save_output(
    first_save, save_doc, output_file, output_dir, user_says, norskGPT_says
):
    output_path = os.path.join(output_dir, output_file)
    content = user_says + "\n" + norskGPT_says + "\n\n"

    if first_save:
        if os.path.exists(output_path):
            # Prompt the user to confirm if they want to delete the existing file
            print(
                "\n\n\nnorskGPT:  En output-fil finnes allerede. Vil du gi den et nytt navn? (y/n):"
            )
            response = input("User: ")
            if response.lower() == "y":
                print("\nnorskGPT: Vennligst skriv inn et nytt filnavn: ")
                new_file_name = input("User: ")
                if save_doc:
                    new_file_name = new_file_name.strip() + ".docx"
                else:
                    new_file_name = new_file_name.strip() + ".txt"
                new_output_path = os.path.join(output_dir, new_file_name)
                os.rename(output_path, new_output_path)
                print(
                    f"\nnorskGPT: Filnavn '{output_file}' endret til '{new_file_name}'.\n\n\n"
                )

        elif not os.path.exists(output_dir):
            # Create the outputs directory
            os.makedirs(output_dir)

        if save_doc:
            document = Document()
            paragraph = document.add_paragraph(content)
            document.save(output_path)
        else:
            with open(output_path, "w") as f:
                f.write(content)
    else:
        # Use the existing output.docx file
        if save_doc:
            document = Document(output_path)
            paragraph = document.add_paragraph(content)
            document.save(output_path)
        else:
            with open(output_path, "a") as f:
                f.write(content)


def main(api_key, mod, max_tokens, temp, output_dir, message):
    openai.api_key = api_key
    session_tokens = 0
    tot_tokens = 0
    save_flag = False
    save_doc = False
    first_save = True

    # Check if user wants to save last response to file
    user_response = input(
        "norskGPT: Vil du at jeg skal lagre som 1. txt, 2. docx ellers vil jeg ikke lagre chatten (1/2)?"
    )
    if int(user_response) == 1:
        output_file = "output.txt"
        save_flag = True
        print("norskGPT: Lagrer som TXT-fil")
    elif int(user_response) == 2:
        output_file = "output.docx"
        save_flag = True
        save_doc = True
        print("norskGPT: Lagrer som DOCX-fil")
    else:
        print("norskGPT: Fortsetter uten å lagre fil.")

    init_message = "\n\n\nnorskGPT: Du kan skrive quit() på hvilken som helst forespørsel om brukerprompt for å avslutte chatten."
    print(init_message)
    while True:
        # Saving conversation if requested to save.
        # if save_flag:
        #    save_output(first_save, save_doc, output_file, output_dir, init_message)
        #    first_save = False

        # Ask for user input
        user_input = input("\n\n\nUser: ")

        # Check if user wants to quit
        if user_input.lower() == "quit()":
            break

        # Combine the system role behavior and user prompt
        message.append({"role": "user", "content": user_input})

        # Generate a completion using the prompt
        response = openai.ChatCompletion.create(
            model=mod, messages=message, temperature=temp
        )

        # Extract the generated response
        generated_text = response["choices"][0]["message"]["content"]

        # Saving conversation if requested to save.
        if save_flag:
            user_says = "User: " + user_input
            norskGPT_says = "norskGPT: " + generated_text
            save_output(
                first_save, save_doc, output_file, output_dir, user_says, norskGPT_says
            )
            first_save = False

        # Print the generated response
        print(f"norskGPT: {generated_text}")

        # Check for tokens used so far
        session_tokens += response["usage"]["total_tokens"]
        tot_tokens += response["usage"]["total_tokens"]
        if tot_tokens > max_tokens:
            print(
                f"\n\nnorskGPT: {session_tokens} antall token brukt i denne økten og {tot_tokens} token brukt siden forrige sjekk. Ønsker du å fortsette samtalen vår (y/n)?"
            )
            user_response = input("\nUser: ")
            if user_response.lower() == "n":
                break
            else:
                tot_tokens = 0
                continue

        # Add the generated response to conversation history
        messages.append({"role": "assistant", "content": generated_text})


if __name__ == "__main__":
    # Set up your OpenAI API credentials

    # Read the API key and system role behavior from the config file
    with open("./config/config.yml", "r") as f:
        config = yaml.safe_load(f)
        api_key = config["api_key"]
        system_role_behavior = config["system_role_behavior"]
        mod = config["model"]
        max_tokens = int(config["max_tokens"])
        temp = int(config["temperature"])
        output_dir = config["output_dir"]

    messages = [{"role": "system", "content": system_role_behavior}]

    main(api_key, mod, max_tokens, temp, output_dir, messages)
